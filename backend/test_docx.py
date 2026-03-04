from docx import Document

doc = Document('DS160exampleForm.docx')
found = False

# check paragraphs
for p in doc.paragraphs:
    if "GIVEN NAME" in p.text:
        print(f"Found in paragraph: '{p.text}'")
        found = True

# check tables
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            if "GIVEN NAME" in cell.text:
                print(f"Found in table cell: '{cell.text}'")
                found = True

if not found:
    print("Could not find GIVEN NAME in standard paragraphs or tables (probably in floating shapes/textboxes).")
